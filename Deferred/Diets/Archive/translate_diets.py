#!/usr/bin/env python3
"""
translate_diets.py — Перевод Diets_v5.md на EN и KA через DeepSeek.
Сохраняет: Diets_v5_EN.md, Diets_v5_KA.md
Затем конвертирует все три в .docx через python-docx.

Запуск:
  cd ~/Desktop/AIM && source venv/bin/activate
  python3 ~/Desktop/Diets/translate_diets.py
"""

import sys, os, re, time
sys.path.insert(0, os.path.expanduser("~/Desktop/AIM"))

from llm import ask_llm

SRC  = os.path.expanduser("~/Desktop/Diets/Diets_v5.md")
OUT_EN = os.path.expanduser("~/Desktop/Diets/Diets_v5_EN.md")
OUT_KA = os.path.expanduser("~/Desktop/Diets/Diets_v5_KA.md")
LOG    = os.path.expanduser("~/Desktop/Diets/translation_log.txt")

CHUNK_LINES = 120   # ~3-4KB per chunk — safe margin for context

# ── Translation prompts ──────────────────────────────────────────────────────

SYSTEM_EN = (
    "You are a professional medical translator specializing in nutrition and dietology. "
    "Translate the following Russian medical text to English. "
    "Preserve all markdown formatting exactly: headers (#, ##, ###), bold (**text**), "
    "italic (*text*), tables (| col |), blockquotes (> text), bullet lists (- item). "
    "Georgian food terms (erba, lobio, tkemali, sulguni, tklapi, gomi, kotsi, tonis puri, "
    "bajhe, jinjoli, ekala) — keep in italic with English explanation in parentheses on first mention. "
    "Georgian names (Vasiko Tkemaladze, Lali Gegeshidze, Jaba Tkemaladze) — keep unchanged. "
    "Medical terms: keep standard English clinical terminology. "
    "Do NOT add any commentary — output only the translated text."
)

SYSTEM_KA = (
    "შენ ხარ პროფესიონალი სამედიცინო მთარგმნელი, სპეციალისტი კვებასა და დიეტოლოგიაში. "
    "თარგმნე შემდეგი რუსული სამედიცინო ტექსტი ქართულად. "
    "შეინარჩუნე markdown-ის ფორმატირება: სათაურები (#, ##, ###), გამუქება (**ტექსტი**), "
    "დახრილი (*ტექსტი*), ცხრილები (| col |), ციტატები (> ტექსტი), სიები (- ელემენტი). "
    "ქართული საკვები ტერმინები (ერბო, ლობიო, ტყემალი, სულუგუნი, ტყლაპი, გომი, "
    "კოცი, თონის პური, ბაჟე, ჯონჯოლი, ეკალა) — გამოიყენე ქართული დასახელება. "
    "სახელები (ვასიკო თქემალაძე, ლალი გეგეშიძე, ჯაბა თქემალაძე) — უცვლელად. "
    "სამედიცინო ტერმინები: გამოიყენე სტანდარტული ქართული სამედიცინო ტერმინოლოგია. "
    "არ დაამატო კომენტარი — გამოიტანე მხოლოდ თარგმნილი ტექსტი."
)

# ── Chunking ─────────────────────────────────────────────────────────────────

def split_into_chunks(text: str, max_lines: int = CHUNK_LINES) -> list[str]:
    """Split on blank lines near chapter/section boundaries."""
    lines = text.split("\n")
    chunks, current, count = [], [], 0
    for line in lines:
        current.append(line)
        count += 1
        if count >= max_lines and (line.strip() == "" or line.startswith("#")):
            chunks.append("\n".join(current))
            current, count = [], 0
    if current:
        chunks.append("\n".join(current))
    return chunks

# ── Translation ──────────────────────────────────────────────────────────────

def translate_chunk(chunk: str, system: str, lang: str, idx: int, total: int) -> str:
    print(f"  [{idx+1}/{total}] {lang} — {len(chunk)} chars...", end=" ", flush=True)
    result = ask_llm(
        prompt=chunk,
        system=system,
        max_tokens=4096,
        temperature=0.2,
    )
    print("✓")
    time.sleep(0.5)   # rate limit courtesy pause
    return result

def translate_file(src_text: str, system: str, lang: str) -> str:
    chunks = split_into_chunks(src_text)
    print(f"\n── Перевод {lang}: {len(chunks)} чанков ──")
    translated = []
    for i, chunk in enumerate(chunks):
        result = translate_chunk(chunk, system, lang, i, len(chunks))
        translated.append(result)
    return "\n".join(translated)

# ── Quality check ─────────────────────────────────────────────────────────────

def quality_check(original_chunk: str, translated_chunk: str, lang: str) -> str:
    prompt = (
        f"Compare this original Russian text with its {lang} translation. "
        f"List any errors: mistranslations, missing content, wrong medical terms, "
        f"broken markdown formatting. Be concise. If no errors, write 'OK'.\n\n"
        f"ORIGINAL:\n{original_chunk[:1500]}\n\n"
        f"TRANSLATION:\n{translated_chunk[:1500]}"
    )
    return ask_llm(prompt, max_tokens=512, temperature=0.1)

# ── Markdown → DOCX ──────────────────────────────────────────────────────────

def md_to_docx(md_text: str, out_path: str, title: str):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re as _re

    doc = Document()
    doc.core_properties.title = title

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def _apply_inline(para, text: str):
        """Apply **bold** and *italic* within a paragraph."""
        pattern = _re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                para.add_run(text[pos:m.start()])
            if m.group(1).startswith("**"):
                run = para.add_run(m.group(2))
                run.bold = True
            else:
                run = para.add_run(m.group(3))
                run.italic = True
            pos = m.end()
        if pos < len(text):
            para.add_run(text[pos:])

    def _table_row(row_text: str) -> list[str]:
        cells = [c.strip() for c in row_text.strip("|").split("|")]
        return cells

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- separator
        if _re.match(r'^---+$', line.strip()):
            i += 1; continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1; continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1; continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1; continue
        if line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            i += 1; continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            _apply_inline(p, line[2:].strip())
            i += 1; continue

        # Table: detect block
        if line.startswith("|"):
            # collect table rows
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = lines[i]
                # skip separator rows |---|---|
                if not _re.match(r'^\|[\s\-:|]+\|', row):
                    rows.append(_table_row(row))
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci < cols:
                            cell = tbl.cell(ri, ci)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            _apply_inline(p, cell_text)
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
                doc.add_paragraph("")
            continue

        # Bullet list
        if _re.match(r'^[\*\-\•]\s', line):
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, line[2:].strip())
            i += 1; continue

        # Numbered list
        if _re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style="List Number")
            text = _re.sub(r'^\d+\.\s', '', line)
            _apply_inline(p, text.strip())
            i += 1; continue

        # Empty line
        if not line.strip():
            i += 1; continue

        # Normal paragraph
        p = doc.add_paragraph()
        _apply_inline(p, line.strip())
        i += 1

    doc.save(out_path)
    print(f"  ✓ DOCX: {out_path}")

# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    print(f"Читаю {SRC}...")
    src_text = open(SRC, encoding="utf-8").read()
    print(f"  {len(src_text)} символов, {src_text.count(chr(10))} строк")

    log_lines = []

    # ── EN translation ────────────────────────────────────────────────────────
    print("\n=== АНГЛИЙСКИЙ ПЕРЕВОД ===")
    en_text = translate_file(src_text, SYSTEM_EN, "EN")
    with open(OUT_EN, "w", encoding="utf-8") as f:
        f.write(en_text)
    print(f"  Сохранён: {OUT_EN}")

    # Quality check EN (first chunk only)
    chunks_src = split_into_chunks(src_text)
    chunks_en  = split_into_chunks(en_text)
    print("\n── Quality check EN (первый чанк) ──")
    qc_en = quality_check(chunks_src[0], chunks_en[0], "English")
    print(f"  {qc_en[:300]}")
    log_lines.append(f"QC EN chunk 0: {qc_en}")

    # ── KA translation ────────────────────────────────────────────────────────
    print("\n=== ГРУЗИНСКИЙ ПЕРЕВОД ===")
    ka_text = translate_file(src_text, SYSTEM_KA, "KA")
    with open(OUT_KA, "w", encoding="utf-8") as f:
        f.write(ka_text)
    print(f"  Сохранён: {OUT_KA}")

    # Quality check KA (first chunk only)
    chunks_ka = split_into_chunks(ka_text)
    print("\n── Quality check KA (первый чанк) ──")
    qc_ka = quality_check(chunks_src[0], chunks_ka[0], "Georgian")
    print(f"  {qc_ka[:300]}")
    log_lines.append(f"QC KA chunk 0: {qc_ka}")

    # Write log
    with open(LOG, "w", encoding="utf-8") as f:
        f.write("\n".join(log_lines))

    # ── DOCX conversion ───────────────────────────────────────────────────────
    print("\n=== DOCX КОНВЕРТАЦИЯ ===")
    md_to_docx(src_text, os.path.expanduser("~/Desktop/Diets/Diets_v5_RU.docx"),
               "Диеты: синтез систем Ткемаладзе-Гегешидзе и Певзнера")
    md_to_docx(en_text, os.path.expanduser("~/Desktop/Diets/Diets_v5_EN.docx"),
               "Diets: Synthesis of the Tkemaladze-Gegeshidze and Pevzner Systems")
    md_to_docx(ka_text, os.path.expanduser("~/Desktop/Diets/Diets_v5_KA.docx"),
               "დიეტები: თქემალაძე-გეგეშიძისა და პევზნერის სისტემების სინთეზი")

    print("\n✅ Готово!")
    print(f"  RU: Diets_v5_RU.docx")
    print(f"  EN: Diets_v5_EN.docx + Diets_v5_EN.md")
    print(f"  KA: Diets_v5_KA.docx + Diets_v5_KA.md")
    print(f"  Log: translation_log.txt")

if __name__ == "__main__":
    main()
