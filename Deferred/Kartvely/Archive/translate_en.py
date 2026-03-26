#!/usr/bin/env python3
"""
Translate Kartvely.md from Russian to English using DeepSeek API.
Processes in chunks of ~130 lines each.
"""

import os
import sys
import time

# ── API key ────────────────────────────────────────────────────────
def get_api_key():
    env_file = os.path.expanduser("~/.aim_env")
    if os.path.exists(env_file):
        for line in open(env_file):
            line = line.strip()
            if line.startswith("DEEPSEEK_API_KEY="):
                return line.split("=", 1)[1].strip()
    return os.environ.get("DEEPSEEK_API_KEY", "")

API_KEY = get_api_key()
if not API_KEY:
    print("ERROR: DEEPSEEK_API_KEY not found")
    sys.exit(1)

from openai import OpenAI
client = OpenAI(api_key=API_KEY, base_url="https://api.deepseek.com")

SYSTEM_PROMPT = """You are a professional academic translator specializing in Caucasian history, mythology, and archaeology.
Translate the following Russian text to English with these rules:

1. PRESERVE all markdown formatting exactly: # headings, ## headings, **bold**, *italic*, | tables, > blockquotes, --- dividers, bullet lists, numbered lists
2. Use academic register throughout — formal, scholarly English
3. Georgian/Caucasian proper names transliteration (use standard scholarly conventions):
   - Хвамли → Khvamli
   - Амиран → Amiran
   - Дали → Dali
   - Гмерти → Gmerti
   - Армази → Armazi
   - Заден → Zaden
   - Копала → Kopala
   - Иахсари → Iakhsari
   - Пиркуши → Pirkushi
   - Мцхета → Mtskheta
   - Колхида → Colchis (use classical form)
   - Арагви → Aragvi
   - Кура → Kura (or Mtkvari where appropriate)
   - Сванети → Svaneti
   - Самегрело → Samegrelo
   - Хевсурети → Khevsureti
   - Аджария → Adjara
   - Картвели → Kartvelian (as adjective), Kartveli (as noun)
   - ГХТ (Гипотеза Хвамлийского Трансфера) → KHT (Khvamli Transfer Hypothesis)
   - Шпеккан → Spekkhan
   - Ачопинтрэ → Achopintre
   - Кажи → Kazhi
   - Медея → Medea
   - Ээт → Aeetes
   - Гелиос → Helios
   - Дурэнки → Duranki
   - Пирмет → Pirmet
   - Мзе → Mze
   - Мтваре → Mtvare
   - Курша → Kursha
   - Гаци и Гаим → Gatsi and Gaim
   - Айнина и Данина → Ainina and Danina
   - Накалакари → Nakalakari
   - Ткаши-Мапа → Tkashi-Mapa
4. Translate ALL text — do not skip any part
5. Output ONLY the translated text, no explanations, no notes, no "Here is the translation:" prefix
6. Preserve blank lines exactly as in the source
7. For the abbreviation ГХТ use KHT; for КГТ also use KHT
"""

def translate_chunk(text: str, chunk_num: int, total: int) -> str:
    """Translate a chunk of text using DeepSeek API."""
    print(f"  Translating chunk {chunk_num}/{total} ({len(text)} chars)...", end=" ", flush=True)

    for attempt in range(3):
        try:
            resp = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"Translate this Russian text to English:\n\n{text}"}
                ],
                max_tokens=4096,
                temperature=0.2,
            )
            result = resp.choices[0].message.content.strip()
            print(f"OK ({len(result)} chars)")
            return result
        except Exception as e:
            print(f"\n  Attempt {attempt+1} failed: {e}")
            if attempt < 2:
                time.sleep(5)

    print("  FAILED after 3 attempts, returning original")
    return text


def split_into_chunks(lines: list, chunk_size: int = 130) -> list:
    """Split lines into chunks, trying to break at paragraph/section boundaries."""
    chunks = []
    i = 0
    while i < len(lines):
        end = min(i + chunk_size, len(lines))

        # Try to extend to a natural break point (blank line or heading)
        # Look back from end to find a good break
        if end < len(lines):
            # Search backwards for a blank line or heading
            best_break = end
            for j in range(end, max(i + chunk_size - 20, i + 1), -1):
                if j < len(lines) and (lines[j].strip() == '' or lines[j].startswith('#')):
                    best_break = j
                    break
            end = best_break

        chunk_lines = lines[i:end]
        chunks.append(''.join(chunk_lines))
        i = end

    return chunks


def main():
    input_path = "/home/oem/Desktop/Kartvely/Kartvely.md"
    output_path = "/home/oem/Desktop/Kartvely/Kartvely_EN.md"
    status_path = "/tmp/kartvely_en_status.txt"

    print(f"Reading {input_path}...")
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    total_lines = len(lines)
    print(f"Total lines: {total_lines}")

    chunks = split_into_chunks(lines, chunk_size=130)
    total_chunks = len(chunks)
    print(f"Split into {total_chunks} chunks")

    # Check if partial output exists
    translated_chunks = []
    start_chunk = 0

    if os.path.exists(output_path + ".partial"):
        print("Found partial translation, loading...")
        with open(output_path + ".partial", 'r', encoding='utf-8') as f:
            partial = f.read()
        # Count how many chunk markers
        marker_count = partial.count("<!-- CHUNK_END -->")
        if marker_count > 0:
            parts = partial.split("<!-- CHUNK_END -->")
            translated_chunks = [p for p in parts[:-1]]  # exclude last empty
            start_chunk = len(translated_chunks)
            print(f"Resuming from chunk {start_chunk + 1}")

    print(f"\nStarting translation from chunk {start_chunk + 1}/{total_chunks}...")

    for i, chunk in enumerate(chunks[start_chunk:], start=start_chunk):
        translated = translate_chunk(chunk, i + 1, total_chunks)
        translated_chunks.append(translated)

        # Save partial result every 5 chunks
        if (i + 1) % 5 == 0 or i == total_chunks - 1:
            partial_content = "<!-- CHUNK_END -->".join(translated_chunks) + "<!-- CHUNK_END -->"
            with open(output_path + ".partial", 'w', encoding='utf-8') as f:
                f.write(partial_content)
            print(f"  [Checkpoint saved at chunk {i+1}]")

        # Small delay to avoid rate limiting
        time.sleep(0.5)

    # Join all chunks
    print("\nJoining translated chunks...")
    full_translation = "\n".join(translated_chunks)

    # Clean up double blank lines that might have been introduced
    import re
    full_translation = re.sub(r'\n{4,}', '\n\n\n', full_translation)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_translation)

    print(f"Translation saved to {output_path}")

    # Remove partial file
    if os.path.exists(output_path + ".partial"):
        os.remove(output_path + ".partial")

    # Now generate DOCX
    print("\nGenerating DOCX...")
    generate_docx(output_path)

    # Write status
    with open(status_path, 'w', encoding='utf-8') as f:
        f.write(f"OK: Kartvely_EN.md and Kartvely_EN_final.docx generated from {total_chunks} chunks ({total_lines} lines)")

    print(f"\nDone! Status written to {status_path}")


def generate_docx(md_path: str):
    """Generate a .docx from the translated markdown."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Installing...")
        os.system("pip install python-docx")
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    import re

    docx_path = "/home/oem/Desktop/Kartvely/Kartvely_EN_final.docx"

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("KARTVELY")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("The Mountain Fortress of the Caucasus\nA Hypothesis on the Origins of Caucasus Connections\nwith the Civilizations of Sumer, Canaan and Egypt")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author.add_run("\nJaba Tkemaladze\nTbilisi — 2024–2026")
    author_run.font.size = Pt(12)

    doc.add_page_break()

    def add_formatted_paragraph(doc, text):
        """Add a paragraph with inline formatting (*italic*, **bold**)."""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)

        # Parse inline formatting
        # Pattern: **bold**, *italic*, ***bold-italic***
        segments = re.split(r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for seg in segments:
            if seg.startswith('***') and seg.endswith('***'):
                run = para.add_run(seg[3:-3])
                run.bold = True
                run.italic = True
            elif seg.startswith('**') and seg.endswith('**'):
                run = para.add_run(seg[2:-2])
                run.bold = True
            elif seg.startswith('*') and seg.endswith('*') and len(seg) > 2:
                run = para.add_run(seg[1:-1])
                run.italic = True
            else:
                if seg:
                    para.add_run(seg)
        return para

    def add_table_row(doc, row_text):
        """Process a markdown table row."""
        # Tables are complex — just add as formatted text for now
        cells = [c.strip() for c in row_text.strip('|').split('|')]
        para = doc.add_paragraph()
        run = para.add_run(' | '.join(cells))
        run.font.size = Pt(10)
        return para

    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip table separator rows
        if re.match(r'^\|[\s\-:|]+\|', line):
            i += 1
            continue

        # Table row
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            cells = [c.strip() for c in line.strip('|').split('|')]
            table_data.append(cells)
            i += 1
            # Check if next line is also table
            if i < len(lines) and (lines[i].startswith('|') or re.match(r'^\|[\s\-:|]+\|', lines[i])):
                continue
            else:
                # End of table — render it
                if table_data:
                    max_cols = max(len(r) for r in table_data)
                    tbl = doc.add_table(rows=len(table_data), cols=max_cols)
                    tbl.style = 'Table Grid'
                    for ri, row in enumerate(table_data):
                        for ci, cell_text in enumerate(row):
                            if ci < max_cols:
                                cell = tbl.cell(ri, ci)
                                cell.text = cell_text
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                in_table = False
                table_data = []
            continue

        in_table = False

        # Horizontal rule
        if line.strip() == '---' or line.strip() == '***' or line.strip() == '___':
            doc.add_paragraph('─' * 40)
            i += 1
            continue

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # H2
        if line.startswith('## ') and not line.startswith('### '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # H4
        if line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            para = doc.add_paragraph(line[2:].strip())
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.italic = True
            i += 1
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('* ') and not line.startswith('**'):
            text = line[2:].strip()
            # Remove inline formatting for list items (simplified)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            para = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            para = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraph
        clean = line.strip()
        if clean:
            add_formatted_paragraph(doc, clean)

        i += 1

    doc.save(docx_path)
    print(f"DOCX saved to {docx_path}")


if __name__ == "__main__":
    main()
