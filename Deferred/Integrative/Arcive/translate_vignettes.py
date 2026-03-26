#!/usr/bin/env python3
"""
Перевод вигнеток Марины через Claude API.

Вигнетки помечены [VIGNETTE — ПЕРЕВЕСТИ С CLAUDE: ...] в v1.docx файлах.
Прогресс сохраняется после каждой вигнетки — можно прерывать и продолжать.

Запуск:
    ANTHROPIC_API_KEY=sk-ant-... python3 translate_vignettes.py --lang en
    ANTHROPIC_API_KEY=sk-ant-... python3 translate_vignettes.py --lang ka
    ANTHROPIC_API_KEY=sk-ant-... python3 translate_vignettes.py --lang kz

    # Все три в фоне:
    ANTHROPIC_API_KEY=sk-ant-... nohup python3 translate_vignettes.py --lang en > logs/vignettes_en.log 2>&1 &
    ANTHROPIC_API_KEY=sk-ant-... nohup python3 translate_vignettes.py --lang ka > logs/vignettes_ka.log 2>&1 &
    ANTHROPIC_API_KEY=sk-ant-... nohup python3 translate_vignettes.py --lang kz > logs/vignettes_kz.log 2>&1 &

    # Прогресс:
    python3 translate_vignettes.py --lang en --status
"""

import argparse
import openai  # DeepSeek совместим с OpenAI SDK
import json
import os
import re
import sys
import time
from docx import Document
from docx.shared import Pt

# ── Конфигурация ───────────────────────────────────────────────────────────────

WORKDIR = "/home/oem/Desktop/Integrative"

LANG_CONFIG = {
    "en": {
        "input":    "Medicine_of_Generations_v1.docx",
        "output":   "Medicine_of_Generations_v2.docx",
        "progress": "vignettes_progress_en.json",
        "lang_name": "English",
        "style_note": "literary, elegant English prose. Preserve the intimate third-person voice.",
    },
    "ka": {
        "input":    "Medicina_Taobata_v1.docx",
        "output":   "Medicina_Taobata_v2.docx",
        "progress": "vignettes_progress_ka.json",
        "lang_name": "Georgian (ქართული)",
        "style_note": "literary Georgian prose. Character name: მარინა. Preserve warm, introspective tone.",
    },
    "kz": {
        "input":    "Urpaktar_Medicinasy_v1.docx",
        "output":   "Urpaktar_Medicinasy_v2.docx",
        "progress": "vignettes_progress_kz.json",
        "lang_name": "Kazakh (Қазақша)",
        "style_note": "literary Kazakh prose. Character name: Марина. Preserve intimate tone.",
    },
}

VIGNETTE_MARKER = "VIGNETTE — ПЕРЕВЕСТИ С CLAUDE"
MODEL = "deepseek-chat"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"

# Контекст персонажа для промпта
CHARACTER_CONTEXT = """
Marina is a 44-year-old book editor in Moscow. She has Hashimoto's thyroiditis and chronic fatigue.
She has not written her own book in 20 years — afraid of another rejection.
Her body's illness is, in part, a defense mechanism against pursuing what she fears most: writing.

The vignettes follow her arc across 73 chapters:
- Chapters 1–15: She seeks a diagnosis, sees doctors, feels unheard
- Chapters 16–30: She begins to notice the connection between her life and her body
- Chapters 31–50: She makes small changes — food, sleep, walks, therapy
- Chapters 51–73: She slowly finds the courage to write

Style: third-person, intimate, spare. 100–180 words per vignette.
Each vignette opens the chapter before the main medical text begins.
"""

# ── Вспомогательные функции ────────────────────────────────────────────────────

def extract_original_from_marker(text: str) -> str:
    """Извлечь оригинальный русский текст из маркера вигнетки."""
    # Формат: [VIGNETTE — ПЕРЕВЕСТИ С CLAUDE: первые 60 символов...]
    match = re.search(r'VIGNETTE — ПЕРЕВЕСТИ С CLAUDE\]?:?\s*(.+?)\.\.\.', text, re.DOTALL)
    if match:
        return match.group(1)
    return text

def is_vignette_para(para) -> bool:
    return VIGNETTE_MARKER in para.text

def load_progress(path: str) -> dict:
    if os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {}  # {str(para_idx): translated_text}

def save_progress(path: str, progress: dict):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(progress, f, ensure_ascii=False, indent=2)

def set_para_italic(para, text: str):
    """Заменить содержимое параграфа на курсивный текст."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        para.runs[0].italic = True
    else:
        run = para.add_run(text)
        run.italic = True

def translate_vignette(client, original_ru: str, lang_name: str, style_note: str,
                        chapter_hint: str = "") -> str:
    """Перевести одну вигнетку через Claude API."""
    prompt = f"""You are a literary translator. Translate the following Russian vignette into {lang_name}.

CHARACTER CONTEXT:
{CHARACTER_CONTEXT}

TRANSLATION STYLE: {style_note}

{f'CHAPTER CONTEXT: {chapter_hint}' if chapter_hint else ''}

RUSSIAN ORIGINAL:
{original_ru}

Translate the full vignette. Output ONLY the translated text, no explanations, no quotes."""

    max_retries = 5
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                max_tokens=600,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if attempt < max_retries - 1:
                wait = 2 ** attempt  # 1, 2, 4, 8, 16 сек
                print(f"  ⚠ Ошибка API (попытка {attempt+1}/{max_retries}): {e}. Жду {wait}с...")
                time.sleep(wait)
            else:
                print(f"  ✗ Не удалось перевести после {max_retries} попыток: {e}")
                return original_ru  # fallback — оставить оригинал

# ── Статус ─────────────────────────────────────────────────────────────────────

def show_status(lang: str):
    cfg = LANG_CONFIG[lang]
    progress_path = os.path.join(WORKDIR, cfg["progress"])
    input_path    = os.path.join(WORKDIR, cfg["input"])

    progress = load_progress(progress_path)
    doc = Document(input_path)
    vignette_idxs = [i for i, p in enumerate(doc.paragraphs) if is_vignette_para(p)]

    done = sum(1 for idx in vignette_idxs if str(idx) in progress)
    total = len(vignette_idxs)
    pct = done / total * 100 if total else 0

    print(f"\n=== Прогресс вигнеток {lang.upper()} ===")
    print(f"Всего вигнеток: {total}")
    print(f"Переведено: {done}")
    print(f"Осталось: {total - done}")
    print(f"Готовность: {pct:.1f}%")
    if os.path.exists(os.path.join(WORKDIR, cfg["output"])):
        print(f"✅ Финальный файл готов: {cfg['output']}")
    else:
        print(f"⏳ Финальный файл ещё не собран")

# ── Основная логика ────────────────────────────────────────────────────────────

def run_translation(lang: str):
    api_key = os.environ.get("DEEPSEEK_API_KEY")
    if not api_key:
        print("❌ Не установлен DEEPSEEK_API_KEY")
        print("   Запуск: DEEPSEEK_API_KEY=sk-... python3 translate_vignettes.py --lang", lang)
        sys.exit(1)

    cfg = LANG_CONFIG[lang]
    input_path    = os.path.join(WORKDIR, cfg["input"])
    output_path   = os.path.join(WORKDIR, cfg["output"])
    progress_path = os.path.join(WORKDIR, cfg["progress"])

    client = openai.OpenAI(api_key=api_key, base_url=DEEPSEEK_BASE_URL)
    progress = load_progress(progress_path)

    doc = Document(input_path)
    paras = doc.paragraphs

    vignette_idxs = [i for i, p in enumerate(paras) if is_vignette_para(p)]
    total = len(vignette_idxs)
    already_done = sum(1 for idx in vignette_idxs if str(idx) in progress)

    print(f"\n=== Перевод вигнеток {lang.upper()} через Claude API ===")
    print(f"Модель: {MODEL} (DeepSeek)")
    print(f"Всего вигнеток: {total}")
    print(f"Уже переведено: {already_done}")
    print(f"Осталось: {total - already_done}")
    if already_done == total:
        print("✅ Все вигнетки переведены. Собираю финальный документ...")
    else:
        print("Начинаю... (Ctrl+C сохранит прогресс)\n")

    new_count = 0
    try:
        for n, idx in enumerate(vignette_idxs):
            key = str(idx)
            if key in progress:
                continue

            para = paras[idx]
            original_preview = extract_original_from_marker(para.text)

            # Находим ближайший заголовок главы (ищем назад)
            chapter_hint = ""
            for j in range(idx - 1, max(0, idx - 30), -1):
                t = paras[j].text.strip()
                if t.startswith("ГЛАВА") or t.startswith("CHAPTER") or len(t) > 10 and t.isupper():
                    chapter_hint = t[:100]
                    break

            # Нужно прочитать полный оригинальный русский текст из progress_XX.json
            # Маркер содержит только первые 60 символов — берём полный из прогресс-файла основного перевода
            base_progress_path = os.path.join(WORKDIR, f"progress_{lang}.json")
            full_original = original_preview
            if os.path.exists(base_progress_path):
                with open(base_progress_path, encoding="utf-8") as f:
                    base_progress = json.load(f)
                # В base_progress под этим idx хранится маркер с превью
                # Нужен оригинал — читаем из v5.docx
                v5_path = os.path.join(WORKDIR, "Медицина_Поколений_v5.docx")
                if os.path.exists(v5_path):
                    v5_doc = Document(v5_path)
                    if idx < len(v5_doc.paragraphs):
                        full_original = v5_doc.paragraphs[idx].text.strip()

            translated = translate_vignette(
                client, full_original, cfg["lang_name"], cfg["style_note"], chapter_hint
            )

            progress[key] = translated
            new_count += 1

            # Статус каждые 10 вигнеток
            done_total = already_done + new_count
            pct = done_total / total * 100
            print(f"  [{done_total}/{total}] {pct:.1f}% | гл.: {chapter_hint[:40] if chapter_hint else '?'}")

            # Сохраняем прогресс каждые 5 вигнеток
            if new_count % 5 == 0:
                save_progress(progress_path, progress)

            # Небольшая пауза чтобы не превысить rate limit
            time.sleep(0.5)

    except KeyboardInterrupt:
        print("\n⏸ Прервано. Сохраняю прогресс...")
        save_progress(progress_path, progress)
        done_total = already_done + new_count
        print(f"✅ Прогресс сохранён в {progress_path}")
        print(f"   Переведено: {done_total}/{total}")
        print("Запустите снова чтобы продолжить.")
        sys.exit(0)

    save_progress(progress_path, progress)

    # ── Сборка финального документа ──────────────────────────────────────────

    print(f"\nСобираю финальный {cfg['output']}...")
    out_doc = Document(input_path)
    out_paras = out_doc.paragraphs

    replaced = 0
    for idx in vignette_idxs:
        key = str(idx)
        if key in progress and idx < len(out_paras):
            set_para_italic(out_paras[idx], progress[key])
            replaced += 1

    out_doc.save(output_path)

    print(f"✅ Сохранено: {output_path}")
    print(f"\n=== Готово ===")
    print(f"Вигнеток заменено: {replaced}/{total}")
    print(f"Файл готов к редактуре: {cfg['output']}")


# ── Точка входа ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    # Создать папку для логов если нужно
    os.makedirs(os.path.join(WORKDIR, "logs"), exist_ok=True)

    parser = argparse.ArgumentParser(description="Перевод вигнеток Марины через Claude API")
    parser.add_argument("--lang", choices=["en", "ka", "kz"], required=True)
    parser.add_argument("--status", action="store_true")
    args = parser.parse_args()

    os.chdir(WORKDIR)

    if args.status:
        show_status(args.lang)
    else:
        run_translation(args.lang)
