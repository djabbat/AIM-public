#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Правки Медицина_Поколений_v4.docx → v5.docx
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import copy

INPUT  = '/home/oem/Desktop/Integrative/Медицина_Поколений_v4.docx'
OUTPUT = '/home/oem/Desktop/Integrative/Медицина_Поколений_v5.docx'

doc = Document(INPUT)
paras = doc.paragraphs

report = []


# ─────────────────────────────────────────────────────────────
# Вспомогательные функции
# ─────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    """Замена текста в runs, сохраняя форматирование."""
    changed = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed


def delete_para(para):
    p = para._element
    p.getparent().remove(p)


def make_italic_para(para, text):
    """Заменить содержимое параграфа на курсивный текст."""
    # Удалить все существующие runs
    for run in para.runs:
        run.text = ''
    # Очистить XML от лишних runs
    p_elem = para._element
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    # Добавить новый run с курсивом
    run = para.add_run(text)
    run.italic = True


# ─────────────────────────────────────────────────────────────
# ПРАВКА 1: Удалить дублирующий протокол питания из главы 73
# Параграфы 3570–3594 (включительно)
# ─────────────────────────────────────────────────────────────

# Найдём границы блока: от первого ⚠ после параграфа 3565 до 3594
# Параграфы 3570–3594 — это 25 параграфов протокола

# Нужно удалять начиная с конца (чтобы не смещать индексы)
# Сначала соберём список объектов-параграфов для удаления
delete_start = 3570
delete_end   = 3594  # включительно

# Убедимся, что параграф 3570 содержит '⚠'
p_check = paras[delete_start]
assert '⚠' in p_check.text, f"Ожидали ⚠ в параграфе {delete_start}, получили: {p_check.text[:80]!r}"
p_check_end = paras[delete_end]

# Собираем объекты (до удаления, чтобы индексы не смещались)
to_delete = [paras[i] for i in range(delete_start, delete_end + 1)]

for p in to_delete:
    delete_para(p)

report.append(f"Правка 1: удалены параграфы {delete_start}–{delete_end} ({len(to_delete)} шт.) — протокол питания в главе 73")

# После удаления нужно обновить список параграфов
paras = doc.paragraphs


# ─────────────────────────────────────────────────────────────
# ПРАВКА 2: Исправить сломанное предложение в главе 73
# ─────────────────────────────────────────────────────────────

fix2_old = 'стремясь его увидеть спрашивает'
fix2_new = 'целостно и спрашивает:'
fix2_prefix = 'Интегративная медицина смотрит на человека, стремясь его увидеть спрашивает:'
fix2_target = 'Интегративная медицина смотрит на человека целостно и спрашивает:'

fixed2 = 0
for para in paras:
    if fix2_old in para.text:
        changed = replace_in_para(para, 'смотрит на человека, стремясь его увидеть спрашивает:', 'смотрит на человека целостно и спрашивает:')
        if changed:
            fixed2 += 1
        break

report.append(f"Правка 2: исправлено сломанных предложений: {fixed2}")


# ─────────────────────────────────────────────────────────────
# ПРАВКА 3: Унифицировать терминологию PNI/PNIE/ПНИЕ
# ─────────────────────────────────────────────────────────────

count_pni  = 0  # PNI → PNIE (где нет уже PNIE)
count_pnie_ru = 0  # ПНИЕ → PNIE
count_doublen = 0  # двойная н

for para in paras:
    # психонейроиммунноэндокринология → психонейроиммуноэндокринология
    if 'психонейроиммунноэндокринология' in para.text:
        if replace_in_para(para, 'психонейроиммунноэндокринология', 'психонейроиммуноэндокринология'):
            count_doublen += 1

    # ПНИЕ → PNIE
    if 'ПНИЕ' in para.text:
        if replace_in_para(para, 'ПНИЕ', 'PNIE'):
            count_pnie_ru += 1

    # PNI → PNIE, но только если не уже PNIE
    # Нужно заменить 'PNI' только тогда, когда следующий символ не 'E'
    if 'PNI' in para.text:
        # Аккуратная замена: заменяем PNI только если следом не E
        for run in para.runs:
            if 'PNI' in run.text:
                new_text = ''
                i = 0
                modified = False
                while i < len(run.text):
                    if run.text[i:i+4] == 'PNIE':
                        new_text += 'PNIE'
                        i += 4
                    elif run.text[i:i+3] == 'PNI':
                        new_text += 'PNIE'
                        i += 3
                        modified = True
                    else:
                        new_text += run.text[i]
                        i += 1
                if modified:
                    run.text = new_text
                    count_pni += 1

report.append(f"Правка 3: PNI→PNIE: {count_pni} замен; ПНИЕ→PNIE: {count_pnie_ru}; двойная-н: {count_doublen}")


# ─────────────────────────────────────────────────────────────
# ПРАВКА 4: Исправить грамматические ошибки в вигнетках главы 37
# + заменить блок с медицинскими дозировками на художественную сцену
# ─────────────────────────────────────────────────────────────

# Находим главу 37 (заново, после удаления параграфов)
ch37_start = None
ch37_end   = len(paras)
for i, p in enumerate(paras):
    if p.style.name.startswith('Heading 1') and '37' in p.text and 'ФИТОТЕРАПИЯ' in p.text:
        ch37_start = i
    elif ch37_start and p.style.name.startswith('Heading 1') and i > ch37_start:
        ch37_end = i
        break

assert ch37_start is not None, "Глава 37 не найдена"

MARINA_VIGNETTE = (
    'Марина готовит завтрак. Орехи, что-то зелёное, стакан воды. '
    'Год назад она не задумывалась об этом — ела на бегу, между правками рукописей. '
    'Теперь это первые десять минут дня, которые принадлежат только ей. '
    'Терапевт сказал: тело учится доверять, когда видит регулярность. '
    'Марина не уверена, что верит в это. '
    'Но она замечает: в те утра, когда завтрак есть, первая половина дня проходит иначе. '
    'Тише, что ли.'
)

# Признак блока с дозировками
DOSAGE_MARKERS = ['1-3 столовые', 'столовые ложке', '6-9', 'минут до еды', 'принимать 20-40']

# Найдём курсивные параграфы в главе 37
italic_paras_37 = []
for i in range(ch37_start, ch37_end):
    p = paras[i]
    is_italic = any(run.italic for run in p.runs if run.text.strip())
    if is_italic:
        italic_paras_37.append((i, p))

print(f"Italic paragraphs in ch37: {[(i, p.text[:60]) for i,p in italic_paras_37]}")

# Проверяем: есть ли дозировки в вигнетке
has_dosage = any(
    any(marker in p.text for marker in DOSAGE_MARKERS)
    for _, p in italic_paras_37
)

grammar_fixes = 0
vignette_replaced = False

if has_dosage and italic_paras_37:
    # Заменяем первый курсивный параграф на художественную сцену
    first_idx, first_para = italic_paras_37[0]
    make_italic_para(first_para, MARINA_VIGNETTE)
    vignette_replaced = True

    # Удаляем остальные курсивные параграфы вигнетки (кроме первого)
    # Удаляем с конца, чтобы не смещать
    for _, p in reversed(italic_paras_37[1:]):
        delete_para(p)

    report.append(f"Правка 4: вигнетка главы 37 заменена на художественную сцену (было {len(italic_paras_37)} italic параграфов)")
    # Обновляем список
    paras = doc.paragraphs
else:
    # Только грамматические правки
    grammar_map = [
        ('столовые ложке', 'столовые ложки'),
        ('снижение уровень кортизола', 'снижение уровня кортизола'),
        ('заправлять еду зеленья', 'заправлять еду зеленью'),
        ('этого не достаточно', 'этого недостаточно'),
    ]
    for i, p in italic_paras_37:
        for old, new in grammar_map:
            if old in p.text:
                if replace_in_para(p, old, new):
                    grammar_fixes += 1
        # Дефис в начале прямой речи → тире
        for run in p.runs:
            if run.text.startswith('- ') or run.text.startswith('\n- '):
                run.text = run.text.replace('- ', '— ', 1)
                grammar_fixes += 1
    report.append(f"Правка 4: грамматических исправлений в вигнетке главы 37: {grammar_fixes}")


# ─────────────────────────────────────────────────────────────
# ПРАВКА 5: Расширить вигнетки в главах 64, 66, 69
# ─────────────────────────────────────────────────────────────

VIGNETTE_64 = (
    'Марина нашла в архиве мамину тетрадь. Записи 1987 года — упражнения из санатория, '
    'куда маму отправили с "астеническим синдромом". Почерк усталый, буквы съезжают вниз '
    'к середине страницы. "День 3. Плавание 20 мин. Лучше". "День 7. Боль в спине. Не вышла". '
    '"День 12. Домой".'
)
VIGNETTE_64_2 = (
    'Марина закрыла тетрадь и долго смотрела в окно. Тело не просто болеет. '
    'Тело ведёт записи. Ведёт их точнее, чем мы сами.'
)
VIGNETTE_64_3 = (
    'На следующее утро она вышла на прогулку раньше обычного. Просто так. Без плана.'
)

VIGNETTE_66 = (
    'В три часа ночи Марина не спит. Не из-за тревоги — из-за идеи. '
    'Она только что поняла, как начать первую главу своей книги. '
    'Не ту, которую боялась написать двадцать лет. Просто — первую главу. Одну.'
)
VIGNETTE_66_2 = (
    'Она не встаёт. Она знает: если встанет — мозг включится, и идея растворится в списке дел.'
)
VIGNETTE_66_3 = (
    'Она повторяет про себя первое предложение. Снова и снова. Пока не засыпает.'
)
VIGNETTE_66_4 = 'Утром оно всё ещё там.'

VIGNETTE_69 = (
    'Марина заметила: она перестала извиняться за паузы в разговоре.'
)
VIGNETTE_69_2 = (
    'Раньше — любое молчание надо было заполнить. Сказать что-нибудь, уточнить, объяснить. '
    'Теперь она иногда просто молчит. И видит, что собеседник тоже начинает думать.'
)
VIGNETTE_69_3 = (
    'Это мелочь. Но она помнит, как терапевт говорил: изменения начинаются с того, '
    'что перестаёшь делать что-то привычное. Не добавляешь новое — убираешь старое.'
)
VIGNETTE_69_4 = (
    'Усталость сегодня — три из десяти. Она записала это в дневник. '
    'Добавила: "Кажется, это моя норма теперь".'
)


def add_italic_para_after(ref_para, text):
    """Вставить новый курсивный параграф после ref_para."""
    from docx.oxml import OxmlElement
    from lxml import etree
    import copy as copy_mod

    # Клонируем параграф ref_para как основу
    new_p = copy_mod.deepcopy(ref_para._element)
    # Очищаем runs в клоне
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # Вставляем после ref_para
    ref_para._element.addnext(new_p)

    # Находим объект параграфа
    # Создаём через docx API: добавим run в new_p напрямую
    from docx.oxml import OxmlElement as OE
    r_elem = OE('w:r')
    rPr = OE('w:rPr')
    i_elem = OE('w:i')
    iCs_elem = OE('w:iCs')
    rPr.append(i_elem)
    rPr.append(iCs_elem)
    r_elem.append(rPr)
    t_elem = OE('w:t')
    t_elem.text = text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    new_p.append(r_elem)

    return new_p


def replace_vignette(ch_paras_idx, new_texts):
    """
    ch_paras_idx: list of (idx, para_obj) of italic paragraphs in chapter
    new_texts: list of strings for new paragraphs
    Replaces first italic para, deletes rest, adds new paras after first.
    """
    if not ch_paras_idx:
        return 0

    first_idx, first_para = ch_paras_idx[0]

    # Если вигнетка уже достаточно длинная (>2 предложений), пропустить
    total_text = ' '.join(p.text for _, p in ch_paras_idx)
    word_count = len(total_text.split())
    if word_count > 40:  # уже расширена
        return 0

    # Заменяем первый параграф
    make_italic_para(first_para, new_texts[0])

    # Удаляем остальные (с конца)
    for _, p in reversed(ch_paras_idx[1:]):
        delete_para(p)

    # Добавляем новые параграфы после первого (в обратном порядке, т.к. addnext)
    for text in reversed(new_texts[1:]):
        add_italic_para_after(first_para, text)

    return 1


# Найти главы заново (после возможных удалений в правке 4)
paras = doc.paragraphs

def find_chapter_italic_paras(heading_keyword, heading_num):
    ch_start = None
    ch_end   = len(paras)
    for i, p in enumerate(paras):
        if p.style.name.startswith('Heading 1') and str(heading_num) in p.text:
            ch_start = i
        elif ch_start is not None and p.style.name.startswith('Heading 1') and i > ch_start:
            ch_end = i
            break
    if ch_start is None:
        return None, None, None
    italic_list = []
    for i in range(ch_start + 1, min(ch_start + 10, ch_end)):
        p = paras[i]
        is_italic = any(run.italic for run in p.runs if run.text.strip())
        if is_italic:
            italic_list.append((i, p))
        elif italic_list:
            # Если уже нашли italic и встретили не-italic — стоп
            break
    return ch_start, ch_end, italic_list

ch64_start, _, italic_64 = find_chapter_italic_paras('СЕТЕВАЯ', 64)
ch66_start, _, italic_66 = find_chapter_italic_paras('ЭВОЛЮЦИОННАЯ', 66)
ch69_start, _, italic_69 = find_chapter_italic_paras('ФАРМАКОЛОГИЯ', 69)

r64 = replace_vignette(italic_64, [VIGNETTE_64, VIGNETTE_64_2, VIGNETTE_64_3])
r66 = replace_vignette(italic_66, [VIGNETTE_66, VIGNETTE_66_2, VIGNETTE_66_3, VIGNETTE_66_4])
r69 = replace_vignette(italic_69, [VIGNETTE_69, VIGNETTE_69_2, VIGNETTE_69_3, VIGNETTE_69_4])

report.append(f"Правка 5: вигнетки расширены — гл.64: {r64}, гл.66: {r66}, гл.69: {r69}")


# ─────────────────────────────────────────────────────────────
# Сохранить
# ─────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print("=" * 60)
print(f"Сохранено: {OUTPUT}")
print()
for line in report:
    print(" •", line)
print()
print(f"Всего параграфов в итоге: {len(doc.paragraphs)}")
