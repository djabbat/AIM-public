#!/usr/bin/env python3
"""
Перевод «Медицина_Поколений_v5.docx» на EN / KA / KZ.

Стратегия:
  - Обычный текст  → Google Translate
  - Курсив (вигнетки Марины) → помечается [VIGNETTE] и НЕ переводится машинно,
    чтобы потом перевести вручную с Claude

Resume: прогресс сохраняется в progress_<lang>.json после каждого параграфа.
Повторный запуск продолжает с места остановки.

Запуск:
    python3 translate_v5.py --lang en
    python3 translate_v5.py --lang ka
    python3 translate_v5.py --lang kz
    python3 translate_v5.py --lang en --status   # только показать прогресс
"""

import argparse
import json
import os
import re
import time
import sys

from docx import Document
from docx.shared import Pt
from deep_translator import GoogleTranslator

# ── Конфигурация ───────────────────────────────────────────────────────────────

INPUT  = "/home/oem/Desktop/Integrative/Медицина_Поколений_v5.docx"
OUTDIR = "/home/oem/Desktop/Integrative"

LANG_CONFIG = {
    "en": {
        "target":   "en",
        "output":   "Medicine_of_Generations_v1.docx",
        "progress": "progress_en.json",
    },
    "ka": {
        "target":   "ka",
        "output":   "Medicina_Taobata_v1.docx",
        "progress": "progress_ka.json",
    },
    "kz": {
        "target":   "kk",   # Google Translate код казахского
        "output":   "Urpaktar_Medicinasy_v1.docx",
        "progress": "progress_kz.json",
    },
}

VIGNETTE_MARKER = "[VIGNETTE — ПЕРЕВЕСТИ С CLAUDE]"
BATCH_SAVE = 25   # сохранять прогресс каждые N параграфов

# ── Вспомогательные функции ────────────────────────────────────────────────────

def is_italic_para(para):
    """Параграф считается вигнеткой если все непустые runs — курсив."""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.italic for r in runs)

def is_reference_line(text):
    return bool(re.match(r'^[A-Z][a-zA-Z\-]+,\s+[A-Z]', text))

def is_url_or_doi(text):
    t = text.strip()
    return t.startswith("http") or t.startswith("doi:")

def translate_text(translator, text):
    text = text.strip()
    if not text:
        return text
    if is_reference_line(text):
        return text
    if is_url_or_doi(text):
        return text
    try:
        if len(text) <= 4900:
            result = translator.translate(text)
            return result if result else text
        else:
            parts = re.split(r'(?<=[.!?])\s+', text)
            chunks, chunk = [], ""
            for part in parts:
                if len(chunk) + len(part) < 4900:
                    chunk += (" " if chunk else "") + part
                else:
                    if chunk:
                        chunks.append(translator.translate(chunk))
                        time.sleep(0.3)
                    chunk = part
            if chunk:
                chunks.append(translator.translate(chunk))
            return " ".join(chunks)
    except Exception as e:
        print(f"  ⚠ Ошибка перевода: {e}")
        return text

def load_progress(path):
    if os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {}  # {str(idx): translated_text}

def save_progress(path, progress):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(progress, f, ensure_ascii=False, indent=2)

def set_para_italic(para, text):
    """Очистить runs параграфа и заменить одним курсивным run."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        para.runs[0].italic = True
    else:
        run = para.add_run(text)
        run.italic = True

def set_para_normal(para, text):
    """Очистить runs параграфа и заменить одним обычным run."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        para.runs[0].italic = False
    else:
        para.add_run(text)

# ── Основная логика ────────────────────────────────────────────────────────────

def show_status(lang):
    cfg = LANG_CONFIG[lang]
    progress_path = os.path.join(OUTDIR, cfg["progress"])
    progress = load_progress(progress_path)

    doc = Document(INPUT)
    paras = doc.paragraphs
    total = len(paras)
    non_empty = [p for p in paras if p.text.strip()]
    vignettes = [p for p in non_empty if is_italic_para(p)]
    done = len(progress)
    vign_done = sum(1 for v in progress.values() if v == VIGNETTE_MARKER)

    print(f"\n=== Прогресс {lang.upper()} ===")
    print(f"Всего параграфов: {total}")
    print(f"Непустых: {len(non_empty)}")
    print(f"Вигнеток (курсив): {len(vignettes)}")
    print(f"Переведено: {done} / {len(non_empty)}")
    print(f"Вигнеток помечено (ждут Claude): {vign_done}")
    pct = done / len(non_empty) * 100 if non_empty else 0
    print(f"Готовность: {pct:.1f}%")

def run_translation(lang):
    cfg = LANG_CONFIG[lang]
    progress_path = os.path.join(OUTDIR, cfg["progress"])
    output_path   = os.path.join(OUTDIR, cfg["output"])

    progress = load_progress(progress_path)
    translator = GoogleTranslator(source="ru", target=cfg["target"])

    doc = Document(INPUT)
    paras = doc.paragraphs
    non_empty_idxs = [i for i, p in enumerate(paras) if p.text.strip()]
    total = len(non_empty_idxs)
    already_done = len(progress)

    print(f"\n=== Перевод {lang.upper()} ===")
    print(f"Непустых параграфов: {total}")
    print(f"Уже переведено: {already_done}")
    print(f"Осталось: {total - already_done}")
    if already_done == total:
        print("✅ Перевод уже завершён. Собираю финальный документ...")
    else:
        print("Начинаю... Ctrl+C для паузы (прогресс сохранится)\n")

    # Переводим / берём из кэша
    translated = {}  # idx → text
    vignette_count = 0
    new_count = 0

    try:
        for n, idx in enumerate(non_empty_idxs):
            key = str(idx)
            para = paras[idx]
            orig = para.text.strip()

            if key in progress:
                translated[idx] = progress[key]
                continue

            if is_italic_para(para):
                # Вигнетка — помечаем, не переводим машинно
                translated[idx] = VIGNETTE_MARKER
                vignette_count += 1
            else:
                result = translate_text(translator, orig)
                translated[idx] = result
                time.sleep(0.15)  # не перегружать API

            progress[key] = translated[idx]
            new_count += 1

            # Статус каждые 50 параграфов
            if new_count % 50 == 0:
                done_total = already_done + new_count
                pct = done_total / total * 100
                print(f"  [{done_total}/{total}] {pct:.1f}% | вигнеток помечено: {vignette_count}")

            # Сохраняем прогресс каждые BATCH_SAVE
            if new_count % BATCH_SAVE == 0:
                save_progress(progress_path, progress)

    except KeyboardInterrupt:
        print("\n⏸ Прервано. Сохраняю прогресс...")
        save_progress(progress_path, progress)
        print(f"✅ Прогресс сохранён в {progress_path}")
        print(f"   Переведено: {already_done + new_count}/{total}")
        print("Запустите снова, чтобы продолжить.")
        sys.exit(0)

    save_progress(progress_path, progress)

    # Собираем итоговый документ
    print(f"\nСобираю {cfg['output']}...")
    out_doc = Document(INPUT)
    out_paras = out_doc.paragraphs

    for idx, text in translated.items():
        para = out_paras[idx]
        if text == VIGNETTE_MARKER:
            set_para_italic(para, f"[{VIGNETTE_MARKER}: {para.text.strip()[:60]}...]")
        elif is_italic_para(paras[idx]):
            set_para_italic(para, text)
        else:
            set_para_normal(para, text)

    out_doc.save(output_path)
    print(f"✅ Сохранено: {output_path}")

    # Итоговая статистика
    vign_total = sum(1 for v in progress.values() if v == VIGNETTE_MARKER)
    print(f"\n=== Готово ===")
    print(f"Переведено параграфов: {total - vign_total}")
    print(f"Вигнеток для Claude: {vign_total}")
    print(f"Следующий шаг: перевести вигнетки с Claude")
    print(f"  python3 translate_vignettes.py --lang {lang}")

# ── Точка входа ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Перевод Медицина Поколений v5")
    parser.add_argument("--lang", choices=["en", "ka", "kz"], required=True)
    parser.add_argument("--status", action="store_true", help="Только показать прогресс")
    args = parser.parse_args()

    if args.status:
        show_status(args.lang)
    else:
        run_translation(args.lang)
