#!/usr/bin/env python3
"""
Apply peer review fixes to Medicine_of_Generations_v2.docx
Saves result as Medicine_of_Generations_v3.docx
"""

from docx import Document
from docx.oxml.ns import qn
import re
import os

INPUT  = '/home/oem/Desktop/Integrative/Medicine_of_Generations_v2.docx'
OUTPUT = '/home/oem/Desktop/Integrative/Medicine_of_Generations_v3.docx'

doc = Document(INPUT)
paras = doc.paragraphs

applied = []
not_found = []

# ─────────────────────────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────────────────────────

def delete_para(para):
    p = para._element
    p.getparent().remove(p)

def is_italic(para):
    runs = [r for r in para.runs if r.text.strip()]
    return runs and all(r.italic for r in runs)

def replace_in_para(para, old, new):
    """Replace text across all runs in a paragraph, preserving formatting."""
    # Collect full text and run boundaries
    full_text = para.text
    if old not in full_text:
        return False

    # Simple approach: modify each run individually
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Cross-run replacement: rebuild text in first matching run
    # Build combined text tracking which runs hold which chars
    combined = ''
    run_spans = []
    for run in para.runs:
        start = len(combined)
        combined += run.text
        run_spans.append((start, len(combined), run))

    idx = combined.find(old)
    if idx == -1:
        return False

    end_idx = idx + len(old)
    new_combined = combined[:idx] + new + combined[end_idx:]

    # Redistribute text back to runs
    new_pos = 0
    for (start, end, run) in run_spans:
        # Calculate new start/end after replacement
        # Map old positions to new positions
        def map_pos(old_p):
            if old_p <= idx:
                return old_p
            elif old_p <= end_idx:
                return idx + len(new)
            else:
                return old_p - len(old) + len(new)

        new_start = map_pos(start)
        new_end = map_pos(end)
        run.text = new_combined[new_start:new_end]

    return True


# ─────────────────────────────────────────────────────────────────────────────
# FIX 1: Remove [Researched] marker
# ─────────────────────────────────────────────────────────────────────────────
fix1_done = False
for i, para in enumerate(paras):
    if '[Researched]' in para.text:
        for run in para.runs:
            if '[Researched]' in run.text:
                run.text = run.text.replace(' [Researched]', '').replace('[Researched]', '')
        applied.append("Fix 1: Removed [Researched] marker from para %d" % i)
        fix1_done = True
        break

if not fix1_done:
    # Try full-text approach
    for i, para in enumerate(paras):
        if '[Researched]' in para.text:
            replace_in_para(para, ' [Researched]', '')
            replace_in_para(para, '[Researched]', '')
            applied.append("Fix 1: Removed [Researched] marker (cross-run) from para %d" % i)
            fix1_done = True
            break

if not fix1_done:
    not_found.append("Fix 1: [Researched] marker not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 2: Delete duplicate Ch.73 vignette (short block B)
# Paras 3556 and 3557:
#   3556 = "She closes it. Sits for several minutes in the silence."
#   3557 = 'Then she opens a new document on her computer. At the top, she types a single word: "Chapter."'
# Para 3555 = long DeepSeek vignette (KEEP)
# Para 3558 = continuation "And she begins..." (KEEP)
# ─────────────────────────────────────────────────────────────────────────────
fix2_count = 0
to_delete = []
for i, para in enumerate(paras):
    t = para.text
    if is_italic(para) and (
        (t.strip() == 'She closes it. Sits for several minutes in the silence.') or
        ('opens a new document' in t and 'types a single word' in t)
    ):
        to_delete.append((i, para))

# Delete in reverse order to preserve indices
for i, para in reversed(to_delete):
    applied.append("Fix 2: Deleted Ch.73 short vignette para %d: %r" % (i, para.text[:60]))
    delete_para(para)
    fix2_count += 1

if fix2_count == 0:
    not_found.append("Fix 2: Ch.73 short vignette paras not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 3: Typo in Ch. 15 heading — PSYCHONEUROIMUNOLOGY → PSYCHONEUROIMMUNOLOGY
# ─────────────────────────────────────────────────────────────────────────────
fix3_done = False
# Refresh paras after possible deletions
paras = doc.paragraphs
for i, para in enumerate(paras):
    if 'PSYCHONEUROIMUNOLOGY' in para.text:
        for run in para.runs:
            if 'PSYCHONEUROIMUNOLOGY' in run.text:
                run.text = run.text.replace('PSYCHONEUROIMUNOLOGY', 'PSYCHONEUROIMMUNOLOGY')
        applied.append("Fix 3: Fixed PSYCHONEUROIMUNOLOGY typo in para %d" % i)
        fix3_done = True
        break

if not fix3_done:
    not_found.append("Fix 3: PSYCHONEUROIMUNOLOGY not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 4: Remove zero-width spaces (U+200B)
# ─────────────────────────────────────────────────────────────────────────────
zws_count = 0
paras = doc.paragraphs
for para in paras:
    for run in para.runs:
        if '\u200b' in run.text:
            run.text = run.text.replace('\u200b', '')
            zws_count += 1

if zws_count > 0:
    applied.append("Fix 4: Removed zero-width spaces from %d runs" % zws_count)
else:
    not_found.append("Fix 4: No zero-width spaces found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 5: Russian text in Appendix A → English citation
# ─────────────────────────────────────────────────────────────────────────────
fix5_done = False
replacement_text = "Tkemaladze, V., & Gegeshidze, L. (2026). Nutrition Protocol [Clinical protocol, version 09.03.2026]. Tbilisi."
paras = doc.paragraphs
for i, para in enumerate(paras):
    if 'Протокол питания' in para.text or 'Клинический протокол' in para.text:
        # Clear all runs and set text to first run
        if para.runs:
            para.runs[0].text = replacement_text
            for run in para.runs[1:]:
                run.text = ''
        applied.append("Fix 5: Replaced Russian Appendix A citation in para %d" % i)
        fix5_done = True
        break

if not fix5_done:
    not_found.append("Fix 5: Russian Appendix A text not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 6: Egg cooking time — 40–90 minutes → 10–12 minutes
# ─────────────────────────────────────────────────────────────────────────────
fix6_done = False
paras = doc.paragraphs
for i, para in enumerate(paras):
    if 'hard-boiled' in para.text.lower() and '40' in para.text:
        # Replace both dash variants
        for run in para.runs:
            if '40\u201390 minutes' in run.text:
                run.text = run.text.replace('40\u201390 minutes', '10\u201312 minutes')
                fix6_done = True
            elif '40-90 minutes' in run.text:
                run.text = run.text.replace('40-90 minutes', '10-12 minutes')
                fix6_done = True
        if fix6_done:
            applied.append("Fix 6: Fixed egg cooking time in para %d" % i)
            break

if not fix6_done:
    # Try cross-run replacement
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if 'hard-boiled' in para.text.lower() and '40' in para.text:
            if replace_in_para(para, '40\u201390 minutes', '10\u201312 minutes'):
                applied.append("Fix 6: Fixed egg cooking time (cross-run) in para %d" % i)
                fix6_done = True
                break
            elif replace_in_para(para, '40-90 minutes', '10-12 minutes'):
                applied.append("Fix 6: Fixed egg cooking time (cross-run) in para %d" % i)
                fix6_done = True
                break

if not fix6_done:
    not_found.append("Fix 6: hard-boiled + 40-90 minutes not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 7: Chapter 20 heading → UPPERCASE
# ─────────────────────────────────────────────────────────────────────────────
fix7_done = False
paras = doc.paragraphs
for i, para in enumerate(paras):
    if para.style.name.startswith('Heading') and 'Chapter 20.' in para.text:
        new_text = para.text.upper()
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''
        applied.append("Fix 7: Uppercased Chapter 20 heading in para %d: %r" % (i, new_text))
        fix7_done = True
        break

if not fix7_done:
    not_found.append("Fix 7: Chapter 20 heading not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 8: Gendered pronouns She → It (for medicine/book, not Marina)
# ─────────────────────────────────────────────────────────────────────────────
fix8_count = 0
paras = doc.paragraphs

# (a) "She gives questions worth living with" → "It gives questions worth living with"
for i, para in enumerate(paras):
    if 'She gives questions worth living with' in para.text:
        for run in para.runs:
            if 'She gives questions worth living with' in run.text:
                run.text = run.text.replace('She gives questions worth living with',
                                            'It gives questions worth living with')
                fix8_count += 1
                applied.append("Fix 8a: She gives → It gives in para %d" % i)
                break
        else:
            if replace_in_para(para, 'She gives questions worth living with',
                               'It gives questions worth living with'):
                fix8_count += 1
                applied.append("Fix 8a: She gives → It gives (cross-run) in para %d" % i)

# (b) "She is excellent at treating acute conditions" → "It excels at treating acute conditions"
for i, para in enumerate(paras):
    if 'She is excellent at treating acute conditions' in para.text:
        for run in para.runs:
            if 'She is excellent at treating acute conditions' in run.text:
                run.text = run.text.replace('She is excellent at treating acute conditions',
                                            'It excels at treating acute conditions')
                fix8_count += 1
                applied.append("Fix 8b: She is excellent → It excels in para %d" % i)
                break
        else:
            if replace_in_para(para, 'She is excellent at treating acute conditions',
                               'It excels at treating acute conditions'):
                fix8_count += 1
                applied.append("Fix 8b: She is excellent → It excels (cross-run) in para %d" % i)

if fix8_count == 0:
    not_found.append("Fix 8: She gives / She is excellent not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 9: 'stuck' → lock in constitutive activation context
# ─────────────────────────────────────────────────────────────────────────────
fix9_done = False
paras = doc.paragraphs
for i, para in enumerate(paras):
    if ('constitutive activation' in para.text or 'oncogenic mutations' in para.text) and '"stuck"' in para.text:
        for run in para.runs:
            if '"stuck"' in run.text:
                run.text = run.text.replace('"stuck"', 'lock')
                fix9_done = True
            elif '\u201cstuck\u201d' in run.text:
                run.text = run.text.replace('\u201cstuck\u201d', 'lock')
                fix9_done = True
        if fix9_done:
            applied.append("Fix 9: Replaced 'stuck' with lock in para %d" % i)
            break

if not fix9_done:
    # Try with curly quotes
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if ('constitutive activation' in para.text or 'oncogenic mutations' in para.text):
            if '\u201cstuck\u201d' in para.text:
                if replace_in_para(para, '\u201cstuck\u201d', 'lock'):
                    applied.append("Fix 9: Replaced \u201cstuck\u201d with lock in para %d" % i)
                    fix9_done = True
                    break
            elif '"stuck"' in para.text:
                if replace_in_para(para, '"stuck"', 'lock'):
                    applied.append("Fix 9: Replaced \"stuck\" with lock in para %d" % i)
                    fix9_done = True
                    break

if not fix9_done:
    not_found.append("Fix 9: 'stuck' + constitutive activation not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 10: Standardize Hashimoto's (all occurrences already have apostrophe — verify)
# ─────────────────────────────────────────────────────────────────────────────
fix10_count = 0
paras = doc.paragraphs
for i, para in enumerate(paras):
    # Check for "Hashimoto thyroiditis" without apostrophe
    if re.search(r"Hashimoto(?!['\u2019])\s+thyroiditis", para.text):
        for run in para.runs:
            if re.search(r"Hashimoto(?!['\u2019])\s+thyroiditis", run.text):
                run.text = re.sub(r"Hashimoto(?!['\u2019])(\s+thyroiditis)",
                                  r"Hashimoto's\1", run.text)
                fix10_count += 1

if fix10_count > 0:
    applied.append("Fix 10: Standardized Hashimoto's in %d runs" % fix10_count)
else:
    not_found.append("Fix 10: No 'Hashimoto thyroiditis' (without apostrophe) found — all already correct")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 11: Para 610 — fix broken sentence
# ─────────────────────────────────────────────────────────────────────────────
fix11_done = False
paras = doc.paragraphs
NEW_TEXT_11 = ("Psychological interventions have immunological consequences "
               "\u2014 as does an unconscious desire for an unattainable goal.")
for i, para in enumerate(paras):
    if 'Psychological interventions' in para.text and 'unconscious desire' in para.text:
        # Replace entire paragraph text
        if para.runs:
            full = para.text
            # Keep everything after the fixed sentence (there may be more text)
            # The instruction says to replace the lead sentence; let's replace the whole para
            # Actually keep the rest of the paragraph
            suffix_start = full.find('Psychotherapy, meditation')
            if suffix_start != -1:
                suffix = ' ' + full[suffix_start:]
            else:
                suffix = ''
            new_full = NEW_TEXT_11 + suffix
            para.runs[0].text = new_full
            for run in para.runs[1:]:
                run.text = ''
        applied.append("Fix 11: Fixed para 610 (Psychological interventions) at para %d" % i)
        fix11_done = True
        break

if not fix11_done:
    not_found.append("Fix 11: Psychological interventions + unconscious desire not found")


# ─────────────────────────────────────────────────────────────────────────────
# FIX 12: "gray matter in the brain" → "gray matter"
# (The instruction says "gray matter of the brain" but the actual text uses "in the brain")
# ─────────────────────────────────────────────────────────────────────────────
fix12_count = 0
paras = doc.paragraphs
for i, para in enumerate(paras):
    # Handle both "of the brain" and "in the brain" variants when appearing right after "gray matter"
    if 'gray matter of the brain' in para.text:
        for run in para.runs:
            if 'gray matter of the brain' in run.text:
                run.text = run.text.replace('gray matter of the brain', 'gray matter')
                fix12_count += 1
        if fix12_count == 0:
            if replace_in_para(para, 'gray matter of the brain', 'gray matter'):
                fix12_count += 1
        if fix12_count > 0:
            applied.append("Fix 12: Removed 'of the brain' redundancy in para %d" % i)

    if 'gray matter in the brain' in para.text:
        for run in para.runs:
            if 'gray matter in the brain' in run.text:
                run.text = run.text.replace('gray matter in the brain', 'gray matter')
                fix12_count += 1
        applied.append("Fix 12: Removed 'in the brain' redundancy in para %d" % i)

if fix12_count == 0:
    not_found.append("Fix 12: 'gray matter of/in the brain' not found")


# ─────────────────────────────────────────────────────────────────────────────
# Save document
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
file_size = os.path.getsize(OUTPUT)

# ─────────────────────────────────────────────────────────────────────────────
# Report
# ─────────────────────────────────────────────────────────────────────────────
print("=" * 70)
print("PEER REVIEW FIXES — REPORT")
print("=" * 70)
print(f"\nFile saved: {OUTPUT}")
print(f"File size:  {file_size:,} bytes ({file_size/1024/1024:.2f} MB)")
print(f"\nApplied ({len(applied)} fixes):")
for a in applied:
    print(f"  + {a}")
print(f"\nNot found / already correct ({len(not_found)}):")
for n in not_found:
    print(f"  - {n}")
print("\nDone.")
